"""Generic document processor for multiple file formats"""
import os
import json
import csv
from pathlib import Path
from typing import List, Dict, Any, Optional
import logging

import PyPDF2
import docx
from bs4 import BeautifulSoup
from llama_index.core.schema import Document

from config import DOCUMENT_CONFIG

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Process documents from multiple formats and sources"""
    
    def __init__(self):
        self.supported_formats = DOCUMENT_CONFIG["supported_formats"]
        self.max_file_size = DOCUMENT_CONFIG["max_file_size_mb"] * 1024 * 1024
    
    def load_from_directory(self, directory_path: str) -> List[Document]:
        """Load all supported documents from a directory"""
        documents = []
        directory = Path(directory_path)
        
        if not directory.exists():
            logger.error(f"Directory {directory_path} does not exist")
            return documents
        
        for file_path in directory.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_formats:
                if file_path.stat().st_size > self.max_file_size:
                    logger.warning(f"File {file_path} exceeds size limit, skipping")
                    continue
                
                try:
                    docs = self.process_file(file_path)
                    documents.extend(docs)
                    logger.info(f"✅ Processed {file_path}: {len(docs)} documents")
                except Exception as e:
                    logger.error(f"❌ Error processing {file_path}: {e}")
        
        return documents
    
    def process_file(self, file_path: Path) -> List[Document]:
        """Process a single file based on its format"""
        suffix = file_path.suffix.lower()
        
        if suffix == ".txt" or suffix == ".md":
            return self._process_text_file(file_path)
        elif suffix == ".pdf":
            return self._process_pdf_file(file_path)
        elif suffix == ".docx":
            return self._process_docx_file(file_path)
        elif suffix == ".json":
            return self._process_json_file(file_path)
        elif suffix == ".csv":
            return self._process_csv_file(file_path)
        else:
            logger.warning(f"Unsupported file format: {suffix}")
            return []
    
    def _process_text_file(self, file_path: Path) -> List[Document]:
        """Process text and markdown files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            if not content.strip():
                return []
            
            return [Document(
                text=content,
                metadata={
                    "source": str(file_path),
                    "filename": file_path.name,
                    "format": "text",
                    "size": len(content)
                }
            )]
        except Exception as e:
            logger.error(f"Error reading text file {file_path}: {e}")
            return []
    
    def _process_pdf_file(self, file_path: Path) -> List[Document]:
        """Process PDF files"""
        documents = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        documents.append(Document(
                            text=text,
                            metadata={
                                "source": str(file_path),
                                "filename": file_path.name,
                                "format": "pdf",
                                "page": page_num + 1,
                                "total_pages": len(pdf_reader.pages)
                            }
                        ))
        except Exception as e:
            logger.error(f"Error reading PDF file {file_path}: {e}")
        
        return documents
    
    def _process_docx_file(self, file_path: Path) -> List[Document]:
        """Process Word documents"""
        try:
            doc = docx.Document(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)
            
            if content_parts:
                content = "\n".join(content_parts)
                return [Document(
                    text=content,
                    metadata={
                        "source": str(file_path),
                        "filename": file_path.name,
                        "format": "docx",
                        "paragraphs": len(content_parts)
                    }
                )]
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {e}")
        
        return []
    
    def _process_json_file(self, file_path: Path) -> List[Document]:
        """Process JSON files"""
        documents = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            if isinstance(data, list):
                for i, item in enumerate(data):
                    if isinstance(item, dict):
                        # Try to find text content in common fields
                        text_content = self._extract_text_from_dict(item)
                        if text_content:
                            documents.append(Document(
                                text=text_content,
                                metadata={
                                    "source": str(file_path),
                                    "filename": file_path.name,
                                    "format": "json",
                                    "item_index": i,
                                    **item.get("metadata", {})
                                }
                            ))
                    elif isinstance(item, str) and item.strip():
                        documents.append(Document(
                            text=item,
                            metadata={
                                "source": str(file_path),
                                "filename": file_path.name,
                                "format": "json",
                                "item_index": i
                            }
                        ))
            elif isinstance(data, dict):
                text_content = self._extract_text_from_dict(data)
                if text_content:
                    documents.append(Document(
                        text=text_content,
                        metadata={
                            "source": str(file_path),
                            "filename": file_path.name,
                            "format": "json",
                            **data.get("metadata", {})
                        }
                    ))
        except Exception as e:
            logger.error(f"Error reading JSON file {file_path}: {e}")
        
        return documents
    
    def _process_csv_file(self, file_path: Path) -> List[Document]:
        """Process CSV files"""
        documents = []
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                csv_reader = csv.DictReader(f)
                
                for row_num, row in enumerate(csv_reader):
                    # Combine all row values into text
                    text_parts = []
                    for key, value in row.items():
                        if value and str(value).strip():
                            text_parts.append(f"{key}: {value}")
                    
                    if text_parts:
                        text_content = "\n".join(text_parts)
                        documents.append(Document(
                            text=text_content,
                            metadata={
                                "source": str(file_path),
                                "filename": file_path.name,
                                "format": "csv",
                                "row": row_num + 1,
                                **{k: v for k, v in row.items() if k.startswith("meta_")}
                            }
                        ))
        except Exception as e:
            logger.error(f"Error reading CSV file {file_path}: {e}")
        
        return documents
    
    def _extract_text_from_dict(self, data: Dict[str, Any]) -> str:
        """Extract text content from dictionary"""
        text_fields = ["text", "content", "description", "body", "message", "title", "summary"]
        text_parts = []
        
        for field in text_fields:
            if field in data and data[field]:
                text_parts.append(str(data[field]))
        
        # If no standard text fields, concatenate all string values
        if not text_parts:
            for key, value in data.items():
                if isinstance(value, str) and value.strip() and key != "metadata":
                    text_parts.append(f"{key}: {value}")
        
        return "\n".join(text_parts)
    
    def create_documents_from_text(self, texts: List[str], metadata_list: Optional[List[Dict]] = None) -> List[Document]:
        """Create documents from plain text list"""
        documents = []
        
        for i, text in enumerate(texts):
            if text.strip():
                metadata = metadata_list[i] if metadata_list and i < len(metadata_list) else {}
                metadata.update({
                    "source": "manual_input",
                    "index": i,
                    "format": "text"
                })
                
                documents.append(Document(
                    text=text,
                    metadata=metadata
                ))
        
        return documents